"""
Layout Inference Web Application with Gradio

A Gradio-based layout inference tool that supports image uploads and multiple backend inference engines.
It adopts a reference-style interface design while preserving the original inference logic.
"""

import gradio as gr
import json
import os
import io
import tempfile
import base64
import zipfile
import uuid
import re
import time
from pathlib import Path
from PIL import Image
import requests
import shutil # Import shutil for cleanup
os.environ['PYPANDOC_PANDOC'] = '/usr/local/bin/pandoc'
import pypandoc

# Local tool imports
from dots_ocr.utils import dict_promptmode_to_prompt
from dots_ocr.utils.consts import MIN_PIXELS, MAX_PIXELS
from dots_ocr.utils.demo_utils.display import read_image
from dots_ocr.utils.doc_utils import load_images_from_pdf

# Add DotsOCRParser import
from dots_ocr.parser import DotsOCRParser


# ==================== Configuration ====================
DEFAULT_CONFIG = {
    'ip': "192.168.200.2",
    'port_vllm': 80,
    'min_pixels': MIN_PIXELS,
    'max_pixels': MAX_PIXELS,
    'test_images_dir': "./assets/showcase_origin",
}

# ==================== Global Variables ====================
# Store current configuration
current_config = DEFAULT_CONFIG.copy()

# Create DotsOCRParser instance
dots_parser = DotsOCRParser(
    ip=DEFAULT_CONFIG['ip'],
    port=DEFAULT_CONFIG['port_vllm'],
    dpi=200,
    min_pixels=DEFAULT_CONFIG['min_pixels'],
    max_pixels=DEFAULT_CONFIG['max_pixels']
)

def get_initial_session_state():
    return {
        'processing_results': {
            'original_image': None,
            'processed_image': None,
            'layout_result': None,
            'markdown_content': None,
            'cells_data': None,
            'temp_dir': None,
            'session_id': None,
            'result_paths': None,
            'pdf_results': None
        },
        'pdf_cache': {
            "images": [],
            "current_page": 0,
            "total_pages": 0,
            "file_type": None,
            "is_parsed": False,
            "results": []
        }
    }

def read_image_v2(img):
    """Reads an image, supports URLs and local paths"""
    if isinstance(img, str) and img.startswith(("http://", "https://")):
        with requests.get(img, stream=True) as response:
            response.raise_for_status()
            img = Image.open(io.BytesIO(response.content))
    elif isinstance(img, str):
        img, _, _ = read_image(img, use_native=True)
    elif isinstance(img, Image.Image):
        pass
    else:
        raise ValueError(f"Invalid image type: {type(img)}")
    return img

def load_file_for_preview(file_path, session_state):
    """Loads a file for preview, supports PDF and image files"""
    pdf_cache = session_state['pdf_cache']
    
    if not file_path or not os.path.exists(file_path):
        return None, "<div id='page_info_box'>0 / 0</div>", session_state
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            pages = load_images_from_pdf(file_path)
            pdf_cache["file_type"] = "pdf"
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            image = Image.open(file_path)
            pages = [image]
            pdf_cache["file_type"] = "image"
        else:
            return None, "<div id='page_info_box'>Unsupported file format</div>", session_state
    except Exception as e:
        return None, f"<div id='page_info_box'>PDF loading failed: {str(e)}</div>", session_state
    
    pdf_cache["images"] = pages
    pdf_cache["current_page"] = 0
    pdf_cache["total_pages"] = len(pages)
    pdf_cache["is_parsed"] = False
    pdf_cache["results"] = []
    
    return pages[0], f"<div id='page_info_box'>1 / {len(pages)}</div>", session_state

def turn_page(direction, session_state):
    """Page turning function"""
    pdf_cache = session_state['pdf_cache']
    
    if not pdf_cache["images"]:
        return None, "<div id='page_info_box'>0 / 0</div>", "", session_state

    if direction == "prev":
        pdf_cache["current_page"] = max(0, pdf_cache["current_page"] - 1)
    elif direction == "next":
        pdf_cache["current_page"] = min(pdf_cache["total_pages"] - 1, pdf_cache["current_page"] + 1)

    index = pdf_cache["current_page"]
    current_image = pdf_cache["images"][index]  # Use the original image by default
    page_info = f"<div id='page_info_box'>{index + 1} / {pdf_cache['total_pages']}</div>"
    
    current_json = ""
    if pdf_cache["is_parsed"] and index < len(pdf_cache["results"]):
        result = pdf_cache["results"][index]
        if 'cells_data' in result and result['cells_data']:
            try:
                current_json = json.dumps(result['cells_data'], ensure_ascii=False, indent=2)
            except:
                current_json = str(result.get('cells_data', ''))
        if 'layout_image' in result and result['layout_image']:
            current_image = result['layout_image']
    
    return current_image, page_info, current_json, session_state

def get_test_images():
    """Gets the list of test images"""
    test_images = []
    test_dir = current_config['test_images_dir']
    if os.path.exists(test_dir):
        test_images = [os.path.join(test_dir, name) for name in os.listdir(test_dir) 
                      if name.lower().endswith(('.png', '.jpg', '.jpeg', '.pdf'))]
    return test_images

def create_temp_session_dir():
    """Creates a unique temporary directory for each processing request"""
    session_id = uuid.uuid4().hex[:8]
    temp_dir = os.path.join(tempfile.gettempdir(), f"dots_ocr_demo_{session_id}")
    os.makedirs(temp_dir, exist_ok=True)
    return temp_dir, session_id

def parse_image_with_high_level_api(parser, image, prompt_mode, fitz_preprocess=False):
    """
    Processes using the high-level API parse_image from DotsOCRParser
    """
    # Create a temporary session directory
    temp_dir, session_id = create_temp_session_dir()
    
    try:
        # Save the PIL Image as a temporary file
        temp_image_path = os.path.join(temp_dir, f"input_{session_id}.png")
        image.save(temp_image_path, "PNG")
        
        # Use the high-level API parse_image
        filename = f"demo_{session_id}"
        results = parser.parse_image(
            input_path=image,
            filename=filename, 
            prompt_mode=prompt_mode,
            save_dir=temp_dir,
            fitz_preprocess=fitz_preprocess
        )
        
        # Parse the results
        if not results:
            raise ValueError("No results returned from parser")
        
        result = results[0]  # parse_image returns a list with a single result
        
        layout_image = None
        if 'layout_image_path' in result and os.path.exists(result['layout_image_path']):
            layout_image = Image.open(result['layout_image_path'])
        
        cells_data = None
        if 'layout_info_path' in result and os.path.exists(result['layout_info_path']):
            with open(result['layout_info_path'], 'r', encoding='utf-8') as f:
                cells_data = json.load(f)
        
        md_content = None
        if 'md_content_path' in result and os.path.exists(result['md_content_path']):
            with open(result['md_content_path'], 'r', encoding='utf-8') as f:
                md_content = f.read()
        
        return {
            'layout_image': layout_image,
            'cells_data': cells_data,
            'md_content': md_content,
            'filtered': result.get('filtered', False),
            'temp_dir': temp_dir,
            'session_id': session_id,
            'result_paths': result,
            'input_width': result.get('input_width', 0),
            'input_height': result.get('input_height', 0),
            'processing_time': result.get('processing_time', 0),
        }
    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
        raise e

def parse_pdf_with_high_level_api(parser, pdf_path, prompt_mode):
    """
    Processes using the high-level API parse_pdf from DotsOCRParser
    """
    # Create a temporary session directory
    temp_dir, session_id = create_temp_session_dir()
    
    try:
        # Use the high-level API parse_pdf
        filename = f"demo_{session_id}"
        results = parser.parse_pdf(
            input_path=pdf_path,
            filename=filename,
            prompt_mode=prompt_mode,
            save_dir=temp_dir
        )
        
        # Parse the results
        if not results:
            raise ValueError("No results returned from parser")
        
        # Handle multi-page results
        parsed_results = []
        all_md_content = []
        all_cells_data = []
        
        for i, result in enumerate(results):
            page_result = {
                'page_no': result.get('page_no', i),
                'layout_image': None,
                'cells_data': None,
                'md_content': None,
                'filtered': False,
                'processing_time': result.get('processing_time', i)
            }
            
            # Read the layout image
            if 'layout_image_path' in result and os.path.exists(result['layout_image_path']):
                page_result['layout_image'] = Image.open(result['layout_image_path'])
            
            # Read the JSON data
            if 'layout_info_path' in result and os.path.exists(result['layout_info_path']):
                with open(result['layout_info_path'], 'r', encoding='utf-8') as f:
                    page_result['cells_data'] = json.load(f)
                    all_cells_data.extend(page_result['cells_data'])
            
            # Read the Markdown content
            if 'md_content_path' in result and os.path.exists(result['md_content_path']):
                with open(result['md_content_path'], 'r', encoding='utf-8') as f:
                    page_content = f.read()
                    page_result['md_content'] = page_content
                    all_md_content.append(page_content)
            page_result['filtered'] = result.get('filtered', False)
            parsed_results.append(page_result)
        
        combined_md = "\n\n---\n\n".join(all_md_content) if all_md_content else ""
        return {
            'parsed_results': parsed_results,
            'combined_md_content': combined_md,
            'combined_cells_data': all_cells_data,
            'temp_dir': temp_dir,
            'session_id': session_id,
            'total_pages': len(results)
        }
        
    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
        raise e

# ==================== Core Processing Function ====================
def process_image_inference(session_state, file_input, #test_image_input, file_input,
                          prompt_mode, server_ip, server_port, min_pixels, max_pixels,
                          fitz_preprocess=False
                          ):
    """Core function to handle image/PDF inference"""
    # Use session_state instead of global variables
    processing_results = session_state['processing_results']
    pdf_cache = session_state['pdf_cache']
    
    if processing_results.get('temp_dir') and os.path.exists(processing_results['temp_dir']):
        try:
            shutil.rmtree(processing_results['temp_dir'], ignore_errors=True)
        except Exception as e:
            print(f"Failed to clean up previous temporary directory: {e}")
    
    # Reset processing results for the current session
    session_state['processing_results'] = get_initial_session_state()['processing_results']
    processing_results = session_state['processing_results']
    
    current_config.update({
        'ip': server_ip,
        'port_vllm': server_port,
        'min_pixels': min_pixels,
        'max_pixels': max_pixels
    })
    
    # Update parser configuration
    dots_parser.ip = server_ip
    dots_parser.port = server_port
    dots_parser.min_pixels = min_pixels
    dots_parser.max_pixels = max_pixels
    
    input_file_path = file_input #if file_input else test_image_input
    
    if not input_file_path:
        return None, "Please upload image/PDF file or select test image", "", "", gr.update(value=None), None, "", session_state
    
    file_ext = os.path.splitext(input_file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            # MINIMAL CHANGE: The `process_pdf_file` function is now inlined and uses session_state.
            preview_image, page_info, session_state = load_file_for_preview(input_file_path, session_state)
            pdf_result = parse_pdf_with_high_level_api(dots_parser, input_file_path, prompt_mode)
            
            session_state['pdf_cache']["is_parsed"] = True
            session_state['pdf_cache']["results"] = pdf_result['parsed_results']
            
            processing_results.update({
                'markdown_content': pdf_result['combined_md_content'],
                'cells_data': pdf_result['combined_cells_data'],
                'temp_dir': pdf_result['temp_dir'],
                'session_id': pdf_result['session_id'],
                'pdf_results': pdf_result['parsed_results']
            })
            
            total_elements = len(pdf_result['combined_cells_data'])
            
            # 添加处理时间信息
            processing_time = 0
            if pdf_result['parsed_results']:
                # 取第一个结果的时间作为参考（实际应该取总时间）
                print(f"PDF results: {pdf_result} \n===")
                processing_time = pdf_result['parsed_results'][0].get('processing_time', 0)
            
            # 格式化时间显示
            if processing_time > 60:
                time_display = f"{processing_time/60:.1f} 分钟"
            else:
                time_display = f"{processing_time:.1f} 秒"
                
            info_text = f"**PDF文件数据:**\n- 总页数: {pdf_result['total_pages']}\n- 总分块数: {total_elements}\n- 处理时间: {time_display}\n- Session ID: {pdf_result['session_id']}"
            
            current_page_layout_image = preview_image
            current_page_json = ""
            if session_state['pdf_cache']["results"]:
                first_result = session_state['pdf_cache']["results"][0]
                if 'layout_image' in first_result and first_result['layout_image']:
                    current_page_layout_image = first_result['layout_image']
                if first_result.get('cells_data'):
                    try:
                        current_page_json = json.dumps(first_result['cells_data'], ensure_ascii=False, indent=2)
                    except:
                        current_page_json = str(first_result['cells_data'])
            
            #download_zip_path = None
            #if pdf_result['temp_dir']:
            #    download_zip_path = os.path.join(pdf_result['temp_dir'], f"layout_results_{pdf_result['session_id']}.zip")
            #    with zipfile.ZipFile(download_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            #        for root, _, files in os.walk(pdf_result['temp_dir']):
            #            for file in files:
            #                if not file.endswith('.zip'): zipf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), pdf_result['temp_dir']))
            download_docx_path = None
            if pdf_result['combined_md_content']:
                download_docx_path = create_docx_from_markdown(
                    pdf_result['combined_md_content'],
                    pdf_result['session_id'],
                    pdf_result['temp_dir']
                )
            
            return (
                current_page_layout_image, info_text, pdf_result['combined_md_content'] or "No markdown content generated",
                pdf_result['combined_md_content'] or "No markdown content generated",
                gr.update(value=download_docx_path, visible=bool(download_docx_path)), page_info, current_page_json, session_state
            )
        
        else: # Image processing
            image = read_image_v2(input_file_path)
            session_state['pdf_cache'] = get_initial_session_state()['pdf_cache']
            
            original_image = image
            parse_result = parse_image_with_high_level_api(dots_parser, image, prompt_mode, fitz_preprocess)
            
            if parse_result['filtered']:
                 # 添加处理时间信息
                 processing_time = parse_result.get('processing_time', 0)
                 # 格式化时间显示
                 if processing_time > 60:
                     time_display = f"{processing_time/60:.1f} 分钟"
                 else:
                     time_display = f"{processing_time:.1f} 秒"
                     
                 info_text = f"**Image Information:**\n- Original Size: {original_image.width} x {original_image.height}\n- Processing: JSON parsing failed, using cleaned text output\n- Server: {current_config['ip']}:{current_config['port_vllm']}\n- 处理时间: {time_display}\n- Session ID: {parse_result['session_id']}"
                 processing_results.update({
                     'original_image': original_image, 'markdown_content': parse_result['md_content'],
                     'temp_dir': parse_result['temp_dir'], 'session_id': parse_result['session_id'],
                     'result_paths': parse_result['result_paths']
                 })
                 return original_image, info_text, parse_result['md_content'], parse_result['md_content'], gr.update(visible=False), None, "", session_state
            
            md_content_raw = parse_result['md_content'] or "No markdown content generated"
            processing_results.update({
                'original_image': original_image, 'layout_result': parse_result['layout_image'],
                'markdown_content': parse_result['md_content'], 'cells_data': parse_result['cells_data'],
                'temp_dir': parse_result['temp_dir'], 'session_id': parse_result['session_id'],
                'result_paths': parse_result['result_paths']
            })
            
            # 添加处理时间信息
            processing_time = parse_result.get('processing_time', 0)
            # 格式化时间显示
            if processing_time > 60:
                time_display = f"{processing_time/60:.1f} 分钟"
            else:
                time_display = f"{processing_time:.1f} 秒"
                
            num_elements = len(parse_result['cells_data']) if parse_result['cells_data'] else 0
            info_text = f"**Image Information:**\n- Original Size: {original_image.width} x {original_image.height}\n- Model Input Size: {parse_result['input_width']} x {parse_result['input_height']}\n- Server: {current_config['ip']}:{current_config['port_vllm']}\n- Detected {num_elements} layout elements\n- 处理时间: {time_display}\n- Session ID: {parse_result['session_id']}"
            
            current_json = json.dumps(parse_result['cells_data'], ensure_ascii=False, indent=2) if parse_result['cells_data'] else ""
            
            #download_zip_path = None
            #if parse_result['temp_dir']:
            #    download_zip_path = os.path.join(parse_result['temp_dir'], f"layout_results_{parse_result['session_id']}.zip")
            #    with zipfile.ZipFile(download_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            #        for root, _, files in os.walk(parse_result['temp_dir']):
            #            for file in files:
            #                if not file.endswith('.zip'): zipf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), parse_result['temp_dir']))
            download_docx_path = None
            if parse_result['md_content']:
                download_docx_path = create_docx_from_markdown(
                    parse_result['md_content'],
                    parse_result['session_id'],
                    parse_result['temp_dir']
                )
            
            return (
                parse_result['layout_image'], info_text, parse_result['md_content'] or "No markdown content generated",
                md_content_raw, gr.update(value=download_zip_path, visible=bool(download_zip_path)),
                None, current_json, session_state
            )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return None, f"Error during processing: {e}", "", "", gr.update(value=None), None, "", session_state

# MINIMAL CHANGE: Functions now take `session_state` as an argument.
def clear_all_data(session_state):
    """Clears all data"""
    processing_results = session_state['processing_results']
    
    if processing_results.get('temp_dir') and os.path.exists(processing_results['temp_dir']):
        try:
            shutil.rmtree(processing_results['temp_dir'], ignore_errors=True)
        except Exception as e:
            print(f"Failed to clean up temporary directory: {e}")
    
    # Reset the session state by returning a new initial state
    new_session_state = get_initial_session_state()
    
    return (
        None,  # Clear file input
        "",    # Clear test image selection
        None,  # Clear result image
        "正在识别，请稍侯结果...",  # Reset info display
        "## 正在识别，请稍候结果...",  # Reset Markdown display
        "🕐 正在识别，请稍候结果...",    # Clear raw Markdown text
        gr.update(visible=False),  # Hide download button
        "<div id='page_info_box'>0 / 0</div>",  # Reset page info
        "🕐 正在识别，请稍候结果...",     # Clear current page JSON
        new_session_state
    )

def update_prompt_display(prompt_mode):
    """Updates the prompt display content"""
    return dict_promptmode_to_prompt[prompt_mode]

def _convert_html_tables_to_md(md_text: str) -> str:
    """
    把 md_text 里的 <table>...</table> 块转换为 Pandoc 可用的 pipe table，
    其余内容保持不变。
    """
    table_pattern = re.compile(r'<table.*?>.*?</table>', flags=re.DOTALL | re.IGNORECASE)

    def _single_table_to_md(html: str) -> str:
        # 临时 .html 文件让 Pandoc 读
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
            f.write(html)
            tmp = f.name
        try:
            # html -> markdown_strict (pipe table)
            md = pypandoc.convert_file(tmp, to='markdown_strict', format='html')
            return md.strip()
        finally:
            os.unlink(tmp)

    # 逐段替换
    return table_pattern.sub(lambda m: _single_table_to_md(m.group(0)), md_text)

def wrap_html_tables(markdown_content: str) -> str:
    """
    将 markdown_content 中的 <table>...</table> 包成：
    ```{=html}
    <table>...</table>
    ```
    """
    # 正则：捕获 <table ...> 到 </table> 之间的全部内容（非贪婪）
    table_re = re.compile(r'(<table\b[^>]*>.*?</table>)', flags=re.DOTALL | re.IGNORECASE)

    def wrap_one(match: re.Match) -> str:
        html = match.group(1)
        return f'```{{=html}}\n{html}\n```'

    return table_re.sub(wrap_one, markdown_content)

def add_borders(html: str) -> str:
    # 给 table 加外框
    html = re.sub(r'<table\b', r'<table style="border-collapse:collapse;border:1px solid black"', html, flags=re.I)
    # 给所有 th / td 加内框
    #html = re.sub(r'<(/?[th|td])\b', r'<\1 style="border:1px solid black"', html, flags=re.I)
    html = re.sub(r'<(/?(?:th|td))\b', r'<\1 style="border:1px solid black"', html, flags=re.I)
    return html

def _add_table_borders_docx(docx_path: str):
    # 给每张表加边框：优先套用内置样式“Table Grid”，
    # 若样式不存在则直接写入 <w:tblBorders>
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _set_tbl_borders(tbl, color="000000", size=4, val="single"):
        tblPr = tbl._element.tblPr
        # 移除已有 <w:tblBorders>
        for el in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(el)
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), val)
            e.set(qn('w:sz'), str(size))   # 1/8 pt; 4≈0.5pt
            e.set(qn('w:space'), "0")
            e.set(qn('w:color'), color)
            borders.append(e)
        tblPr.append(borders)

    doc = Document(docx_path)
    for tbl in doc.tables:
        try:
            tbl.style = 'Table Grid'  # 大多数 Word/模板自带此样式
        except Exception:
            _set_tbl_borders(tbl)     # 兜底：直接写 OOXML
    doc.save(docx_path)

def process_with_fade_effect(*args):  
    # 先更新图像为淡化状态  
    yield gr.update(elem_classes=["processing"]), "处理中...", "", "", gr.update(visible=False), "", "", args[-1]  
      
    # 调用原始处理函数  
    result = process_image_inference(*args)  
      
    # 返回最终结果，移除淡化效果  
    final_image = gr.update(value=result[0], elem_classes=["completed"])  
    yield final_image, *result[1:]  
  
def create_docx_from_markdown(markdown_content, session_id, temp_dir):  
    """将markdown内容转换为DOCX文件"""  
    docx_path = os.path.join(temp_dir, f"document_{session_id}.docx")  
    

    # 1. 先用 Python 包住所有 <table>...</table>
    markdown_content = wrap_html_tables(markdown_content) #有效，是表格化的关键步骤
    #markdown_content = add_borders(markdown_content) # 没有实际效果
    # 把三个及以上连续的 - 替换成 Word 分页符
    markdown_content = re.sub(
        r'^\s*-{3,}\s*$',          # 整行只有 3 个及以上 -
        '\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n',
        markdown_content,
        flags=re.MULTILINE
    )
    print(f"输出MD: {markdown_content}")
    
    # 获取脚本所在目录，保证 filter 文件能找到
    filter_path = os.path.join(os.path.dirname(__file__), 'keep-html-table.lua')

    # 如果包含 <table>，先转换；否则原样保留
    #if '<table' in markdown_content.lower():
    #    markdown_content = _convert_html_tables_to_md(markdown_content)

    # 使用pypandoc将markdown转换为docx  
    pypandoc.convert_text(  
        markdown_content,   
        'docx',   
        format='markdown+raw_html+markdown_in_html_blocks',  
        outputfile=docx_path,  
        extra_args=[
            f'--lua-filter={filter_path}'
            # , '--reference-doc=template.docx'
        ] #extra_args=['--reference-doc=template.docx']  # 可选：使用模板  
    )

    # 关键点4：后处理，为所有表格加边框线
    _add_table_borders_docx(docx_path)

    return docx_path

# ==================== Gradio Interface ====================
def create_gradio_interface():
    """Creates the Gradio interface"""
    
    # CSS styles, matching the reference style
    css = """

    @font-face {
        font-family: 'IBM Plex Sans';
        font-style: normal;
        font-weight: 400;
        font-stretch: normal;
        src: url('http://ai.byfunds.com/hystatic/ibmplexsans/v22/lzAA.ttf') format('truetype');
    }

    @font-face {
        font-family: 'IBM Plex Mono';
        font-style: normal;
        font-weight: 400;
        src: url('http://ai.byfunds.com/hystatic/ibmplexmono/v19/n5ig.ttf') format('truetype');
    }
    body, .gradio-container {
        font-family: 'IBM Plex Sans', sans-serif !important;
    }
    code, pre, .gradio-code {
        font-family: 'IBM Plex Mono', monospace !important;
    }

    #parse_button {
        background: #FF576D !important; /* !important 确保覆盖主题默认样式 */
        border-color: #FF576D !important;
    }
    /* 鼠标悬停时的颜色 */
    #parse_button:hover {
        background: #F72C49 !important;
        border-color: #F72C49 !important;
    }
    
    #page_info_html {
        display: flex;
        align-items: center;
        justify-content: center;
        height: 100%;
        margin: 0 12px;
    }

    #page_info_box {
        padding: 8px 20px;
        font-size: 16px;
        border: 1px solid #bbb;
        border-radius: 8px;
        background-color: #f8f8f8;
        text-align: center;
        min-width: 80px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }

    #markdown_output {
        min-height: 800px;
        overflow: auto;
    }

    footer {
        visibility: hidden;
    }
    
    #info_box {
        padding: 10px;
        background-color: #f8f9fa;
        border-radius: 8px;
        border: 1px solid #dee2e6;
        margin: 10px 0;
        font-size: 14px;
    }
    
    #result_image {
        border-radius: 8px;
    }
    
    #markdown_tabs {
        height: 100%;
    }

    #result_image.processing {  
        opacity: 0.5;  
        filter: grayscale(50%);  
        transition: opacity 0.3s ease, filter 0.3s ease;  
    }  
  
    #result_image.completed {  
        opacity: 1;  
        filter: none;  
        transition: opacity 0.3s ease, filter 0.3s ease;  
    }  
    """
    theme = gr.themes.Ocean(
        font=["IBM Plex Sans", "sans-serif"],
        font_mono=["IBM Plex Mono", "monospace"]
    )
    with gr.Blocks(theme=theme, css=css, title='dots.ocr') as demo:
        session_state = gr.State(value=get_initial_session_state())
        
        # Title
        gr.HTML("""
            <div style="display: flex; align-items: center; justify-content: center; margin-bottom: 20px;">
                <h1 style="margin: 0; font-size: 2em;">🔍 灵眸·PDF超级解析助手</h1>
            </div>
            <div style="text-align: center; margin-bottom: 10px;">
                <em>对图像/PDF文件进行布局分析和结构化提取</em>
            </div>
        """)
        
        with gr.Row():
            # Left side: Input and Configuration
            with gr.Column(scale=1, elem_id="left-panel"):
                gr.Markdown("### 📥 文件上传")
                file_input = gr.File(
                    label="上传PDF/图像文件", 
                    type="filepath", 
                    file_types=[".pdf", ".jpg", ".jpeg", ".png"],
                )
                
                #test_images = get_test_images()
                #test_image_input = gr.Dropdown(
                #    label="Or Select an Example",
                #    choices=[""] + test_images,
                #    value="",
                #)

                gr.Markdown("### ⚙️ 工作模式")
                prompt_mode = gr.Dropdown(
                    label="选择工作模式",
                    choices=["完全识别", "布局识别", "文字识别"],
                    value="完全识别",
                )
                
                # Display current prompt content
                prompt_display = gr.Textbox(
                    label="工作提示词预览",
                    value=dict_promptmode_to_prompt[list(dict_promptmode_to_prompt.keys())[0]],
                    lines=4,
                    max_lines=8,
                    interactive=False,
                    show_copy_button=True
                )
                
                with gr.Row():
                    process_btn = gr.Button("🔍 开始识别（每页30~60秒）", variant="primary", scale=2, elem_id="parse_button")
                    #clear_btn = gr.Button("🗑️ 重置", variant="secondary", scale=1)
                
                with gr.Accordion("🛠️ 高级设置", open=False):
                    fitz_preprocess = gr.Checkbox(
                        label="启用图像文件“缩适预处理”工序", 
                        value=True,
                        info="以PDF模式处理图像文件(图像->pdf->200dpi大图像). 如果你的原图比较模糊建议启用."
                    )
                    with gr.Row(visible=False):
                        server_ip = gr.Textbox(label="Server IP", value=DEFAULT_CONFIG['ip'])
                        server_port = gr.Number(label="Port", value=DEFAULT_CONFIG['port_vllm'], precision=0)
                    with gr.Row(visible=False):
                        min_pixels = gr.Number(label="Min Pixels", value=DEFAULT_CONFIG['min_pixels'], precision=0)
                        max_pixels = gr.Number(label="Max Pixels", value=DEFAULT_CONFIG['max_pixels'], precision=0)
            # Right side: Result Display
            with gr.Column(scale=6, variant="compact"):
                with gr.Row():
                    # Result Image
                    with gr.Column(scale=3):
                        gr.Markdown("### 👁️ 预览")
                        result_image = gr.Image(
                            label="布局预览",
                            visible=True,
                            height=800,
                            show_label=False
                        )
                        
                        # Page navigation (shown during PDF preview)
                        with gr.Row():
                            prev_btn = gr.Button("⬅ 上一页", size="sm")
                            page_info = gr.HTML(
                                value="<div id='page_info_box'>0 / 0</div>", 
                                elem_id="page_info_html"
                            )
                            next_btn = gr.Button("下一页 ➡", size="sm")
                        
                        # Info Display
                        info_display = gr.Markdown(
                            "等待接收和处理上传文件...",
                            elem_id="info_box"
                        )
                    
                    # Markdown Result
                    with gr.Column(scale=3):
                        gr.Markdown("### ✔️ 识别结果")
                        
                        with gr.Tabs(elem_id="markdown_tabs"):
                            with gr.TabItem("Markdown渲染预览"):
                                md_output = gr.Markdown(
                                    "## 请点击识别按钮...",
                                    max_height=600,
                                    latex_delimiters=[
                                        {"left": "$$", "right": "$$", "display": True},
                                        {"left": "$", "right": "$", "display": False}
                                    ],
                                    show_copy_button=False,
                                    elem_id="markdown_output"
                                )
                            
                            with gr.TabItem("Markdown Raw Text", visible=False):
                                md_raw_output = gr.Textbox(
                                    value="🕐 Waiting for parsing result...",
                                    label="Markdown Raw Text",
                                    max_lines=100,
                                    lines=38,
                                    show_copy_button=True,
                                    elem_id="markdown_output",
                                    show_label=False
                                )
                            
                            with gr.TabItem("Current Page JSON", visible=False):
                                current_page_json = gr.Textbox(
                                    value="🕐 Waiting for parsing result...",
                                    label="Current Page JSON",
                                    max_lines=100,
                                    lines=38,
                                    show_copy_button=True,
                                    elem_id="markdown_output",
                                    show_label=False
                                )
                
                # Download Button
                with gr.Row():
                    download_btn = gr.DownloadButton(
                        "⬇️ 导出DOCX下载文件",
                        visible=False
                    )
        
        # When the prompt mode changes, update the display content
        prompt_mode.change(
            fn=update_prompt_display,
            inputs=prompt_mode,
            outputs=prompt_display,
        )
        
        # Show preview on file upload
        file_input.upload(
            # fn=lambda file_data, state: load_file_for_preview(file_data, state),
            fn=load_file_for_preview,
            inputs=[file_input, session_state],
            outputs=[result_image, page_info, session_state]
        )
        
        # Also handle test image selection
        #test_image_input.change(
            # fn=lambda path, state: load_file_for_preview(path, state),
        #    fn=load_file_for_preview,
        #    inputs=[test_image_input, session_state],
        #    outputs=[result_image, page_info, session_state]
        #)

        prev_btn.click(
            fn=lambda s: turn_page("prev", s),
            inputs=[session_state], 
            outputs=[result_image, page_info, current_page_json, session_state]
        )
        
        next_btn.click(
            fn=lambda s: turn_page("next", s),
            inputs=[session_state], 
            outputs=[result_image, page_info, current_page_json, session_state]
        )
        
        process_btn.click(
            fn=process_image_inference,
            inputs=[
                session_state, file_input, #test_image_input,
                prompt_mode, server_ip, server_port, min_pixels, max_pixels, 
                fitz_preprocess
            ],
            outputs=[
                result_image, info_display, md_output, md_raw_output,
                download_btn, page_info, current_page_json, session_state
            ]
        )
        
        #clear_btn.click(
        #    fn=clear_all_data,
        #    inputs=[session_state],
        #    outputs=[
        #        file_input, #test_image_input,
        #        result_image, info_display, md_output, md_raw_output,
        #        download_btn, page_info, current_page_json, session_state
        #    ]
        #)
    
    return demo

# ==================== Main Program ====================
if __name__ == "__main__":
    import sys
    port = int(sys.argv[1])
    demo = create_gradio_interface()
    demo.queue().launch(
        server_name="0.0.0.0", 
        server_port=port, 
        debug=True,
        root_path="/dotsocr"
    )
